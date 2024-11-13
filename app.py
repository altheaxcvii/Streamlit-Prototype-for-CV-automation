# %%
import streamlit as st
from openai import OpenAI
from docx import Document

# %%
st.title("CV Analysis with GPT")

def parsetable(tableno):
    joined_text = ""
    for row in doc.tables[tableno].rows:
        prev_cell = ""
        row_text = ""

        for cell in row.cells:
            if cell.text != "" and cell.text != prev_cell:
                prev_cell = cell.text
                row_text = row_text + " " + cell.text
        
        row_text = row_text.strip()
        if row_text != "":
            joined_text = joined_text + "\n" + row_text
    return joined_text.strip()

def read_docx_in_order(file_path):
    doc = Document(file_path)
    paragraph_index = 0
    table_index = 0
    full_docx = ""

    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            full_docx = full_docx + parsetable(table_index) + "\n"
            table_index += 1
        elif element.tag.endswith('p'):
            paragraph = doc.paragraphs[paragraph_index]
            if paragraph.text != "":
                full_docx = full_docx + paragraph.text + "\n"
            paragraph_index += 1
    return full_docx

api_key = st.text_input("Enter your ChatGPT API key", type="password")
candidatename = st.text_input("Enter candidate's name")

prompt = f"""
Please provide a 1-page write-up supporting {candidatename}'s appointment. She is female. The recommendation letter should start with: On behalf of Paeds ACP, I hereby present to the Duke-NUS APT Committee the recommendation for the appointment of Dr Lim Yinru at the rank of Clinical Instructor at Duke-NUS. 

Followed by the following information, in formal and professional tone. Do not print the headers.
1.	Candidate’s educational background in chronological order;
2. Candidate's training background in chronological order, summarized
3.	Candidate’s research achievements and accomplishments
4.	Candidate’s participation in educational initiatives and academic service;
5.	Candidate’s current research and/or teaching interests

End the letter with: The Academic Council has assessed {candidatename}’s credentials and found them to be appropriate for the proposed appointment.

The Paeds ACP recommends that the Duke-NUS APT Committee supports the appointment of {candidatename} as Clinical Instructor at Duke-NUS.
"""

if api_key:
    #openai.api_key = api_key
    client = OpenAI(api_key=api_key)
    st.success("API key set!")

    uploaded_file = st.file_uploader("Upload your CV (.docx)", type="docx")

    if uploaded_file is not None:
        doc = Document(uploaded_file)
        cv_text = read_docx_in_order(uploaded_file)
        st.write("Extracted CV Text:")
        st.text(cv_text)

        # Button to run the analysis using the provided API key
        if st.button("Run Analysis"):
            completion = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": cv_text}
                ]
                )

            analysis = completion.choices[0].message.content.strip()
            st.subheader("Generated Resume:")
            st.write(analysis)
else:
    st.info("Please enter your API key to continue.")
                 


